# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# 全局默认字体
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)  # 小四号
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing_rule = None
style.paragraph_format.line_spacing = Twips(20 * 20)  # 20磅 = 20*20 twips
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)

# 页面边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(16)
    run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.line_spacing = Twips(20 * 20)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

def add_heading_7(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(14)
    run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.line_spacing = Twips(20 * 20)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

def add_sub_heading(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.line_spacing = Twips(20 * 20)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)

def add_para(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)  # 小四号
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.line_spacing = Twips(20 * 20)
    p.paragraph_format.first_line_indent = Pt(24)  # 首行缩进两字符
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

# ========== 正文开始 ==========

add_title('第7章 总结与展望')

add_sub_heading('7.1 个人工作详细说明')

add_para('在本项目中，我负责用户管理、膳食管理、护理管理和AI智能服务四个功能模块的设计与开发工作，涵盖后端接口、前端页面和数据库表结构的设计实现。')

add_para('用户管理模块是系统的基础模块，我完成了基于JWT的身份认证流程开发。后端基于Spring Security搭建认证授权体系，用户登录时通过BCrypt加密比对密码，验证通过后签发包含角色信息的JWT令牌。在common模块中编写SecurityConfig配置类，对URL访问权限进行分层定义，同时使用@PreAuthorize注解实现方法级别的角色控制。前端在utils/request.js中封装了Axios请求拦截器，自动在请求头中附加Bearer令牌，并在响应拦截器中统一处理401和403错误。数据库层面，sys_user表包含手机号、密码、姓名、年龄、性别、角色和状态等字段，手机号设置唯一索引防止重复注册，角色字段存储管理员、健康管家、护士和入住老人四种角色标识。前端用户列表页面在moduleConfig.js中配置了各字段，角色使用下拉选择组件，状态使用开关切换形式，方便管理员操作。')

add_para('膳食管理模块我完成了后端MealController的开发和前端日历与列表双视图的实现。数据库设计方面，meal表包含餐食编号、客户ID、早中晚三餐内容、特殊需求、日期和膳食图片等字段，其中餐食编号采用"MEAL+年月日+序号"的格式由后端自动生成。后端继承BaseCrudController获取标准CRUD接口，分页查询通过LEFT JOIN关联客户表展示客户姓名，膳食图片通过阿里云OSS SDK上传并返回访问URL存储。前端实现了日历视图与列表视图的双模式切换，日历视图在每个日期格子中展示膳食摘要，点击弹出详情面板展示三餐详细信息和图片；列表视图支持按客户和日期筛选排序。日历数据采用按月批量加载的方式，减少请求次数提升性能。')

add_para('护理管理模块涵盖护理级别、护理项目和护理记录三个子模块的开发。数据库设计了nursing_level、nursing_item和nursing_record三张表，护理项目通过level_id外键关联护理级别，护理记录通过customer_id和item_id分别关联客户和护理项目，形成三层数据的完整关联。护理级别预置了全护、半护、自护、特护等九个级别，每个级别下配置了具体的护理项目和执行频率。后端三个Controller均继承BaseCrudController，护理记录查询支持按客户、项目和时间范围多条件筛选。前端护理记录页面包含客户、护理项目和护理人员三个下拉框，我在CrudPage组件中增加了数据缓存机制，首次加载后缓存结果，避免每次打开弹窗都重复请求接口，提升了操作响应速度。')

add_para('AI智能服务模块是我投入精力最多的模块，我独立完成了RAG检索增强生成引擎的全链路实现。后端LlmService封装了与MiMo大模型API的对接，chat方法构造OpenAI兼容格式的请求体调用/v1/chat/completions接口，使用Jackson解析响应结果，兼容MiMo推理模型的reasoning_content字段。embeddings方法调用/v1/embeddings接口实现文本向量化，同时实现了基于字符哈希的降级向量方案，保证API异常时RAG流程仍能基本运转。KnowledgeBaseService实现了文本分块功能，采用500字符分块加50字符重叠的滑动窗口算法，保证语义的上下文连贯性。SimpleVectorStore基于ConcurrentHashMap实现内存向量存储，search方法使用余弦相似度算法检索最相关的文档片段。RagService整合以上组件，将检索结果注入系统提示词后调用LLM生成基于知识库的回答，并扩展了健康分析和护理方案推荐功能。前端AiChat页面实现了聊天气泡布局、欢迎区域和快捷提问，打字机效果通过for循环配合slice方法逐帧更新消息内容，每3个字符暂停15毫秒模拟流畅的逐字显示。对话历史通过localStorage持久化，以手机号为键存储消息数组。在容错处理方面，api/ai.js中封装了withMock函数，catch块自动降级返回Mock演示数据，确保AI服务不可用时页面仍能正常展示。')

add_sub_heading('7.2 工作中所遇到的问题及解决方法')

add_para('开发过程中遇到了不少问题。第一个比较棘手的问题是AI对话的打字机效果实现。最初尝试使用SSE做流式返回，但MiMo API的流式响应格式与标准SSE不兼容，解析频繁出错，尝试了多种数据清洗方案都无法稳定处理。最终改为非流式方案，前端通过for循环配合slice方法逐帧更新消息内容，每3个字符触发一次DOM更新并暂停15毫秒，用setTimeout控制显示速度，模拟出了流畅的打字效果。')

add_para('第二个问题是RAG检索结果质量不稳定。初期采用按固定500字符硬切分的方式，经常把一个完整的句子从中间截断，导致检索到的文档片段语义不完整，大模型无法理解参考资料。后来将分块策略改为500字符分块加50字符重叠的滑动窗口算法，相邻块之间有重叠区域，保证每个块的语义连贯性，检索准确率明显提升。')

add_para('第三个问题是护理管理模块的下拉框数据加载。护理记录页面有客户、护理项目、护理人员三个下拉框，每次打开弹窗都重新请求三个接口，而这些接口涉及多表关联查询耗时较长，导致操作卡顿。后来在CrudPage组件中增加了数据缓存机制，首次加载后将结果存入缓存对象，后续打开弹窗直接使用缓存数据，弹窗打开速度从数百毫秒降低到几乎即时响应。')

add_para('第四个问题是AI服务不可用时前端白屏。测试环境中MiMo API偶尔超时或返回错误，前端直接报错无法使用。最终在api/ai.js中封装了withMock函数，将每个AI接口调用包裹在try-catch中，catch块自动降级返回预设的Mock演示数据，包括AI对话的礼貌回复、健康评估的模拟报告等，确保页面始终可正常展示。')

add_sub_heading('7.3 总结与体会')

add_para('通过本次项目的开发，我对前后端分离架构有了从理论到实践的完整理解。从最初对着需求文档无从下手，到逐步理清业务逻辑、设计数据库表结构、定义接口规范、编写前后端代码、联调测试，整个过程让我深刻体会到软件开发不仅仅是写代码，更重要的是对业务的理解和对问题的分析能力。在项目之前，我对Spring Boot、Vue 3、MyBatis-Plus等框架的了解仅停留在课堂层面，通过这次实践才真正掌握了它们在实际项目中的使用方法和最佳实践。')

add_para('在AI智能服务模块的开发中，我接触到了大语言模型API对接、RAG检索增强生成、文本向量化、余弦相似度检索等前沿技术，这些都是课堂上几乎没有涉及的内容，完全靠自己查阅文档和技术博客来学习。从最初不理解什么是Embedding向量，到后来能够独立实现完整的RAG检索引擎，虽然过程中踩了很多坑，但最终看到AI对话能够流畅运行、知识库问答能够准确回答问题时，成就感非常强烈。这也让我认识到，面对技术难题时不能回避，要敢于尝试、善于查阅资料、勤于调试验证。')

add_para('同时，团队协作的过程也让我学到了很多。与负责其他模块的同学沟通接口约定、协调数据格式、联调测试，让我理解了良好的沟通和规范的接口文档对团队开发效率的重要性。项目初期由于接口命名和数据格式不统一，联调浪费了大量时间，后来统一使用Knife4j生成的OpenAPI文档作为规范，效率才明显提升。这次毕业设计是我大学期间最完整的一次工程实践，为我今后的工作打下了坚实的基础。')

# 保存文件
output_path = r'D:\桌面\neusoftelderlycare1\第7章_个人总结.docx'
doc.save(output_path)
print(f'文档已生成: {output_path}')
